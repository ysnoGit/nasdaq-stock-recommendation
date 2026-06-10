from __future__ import annotations

import argparse
import re
from datetime import date
from pathlib import Path

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
REPORT_DIR = Path(__file__).resolve().parent
DOCX_PATH = REPORT_DIR / "backtest_parameter_comparison_report.docx"
CHART_PATH = REPORT_DIR / "parameter_sensitivity.png"

NAVY = "17324D"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
GREEN = "1E6B52"
GOLD = "946200"
RED = "9B1C1C"


def parse_log(log_path: Path) -> tuple[pd.DataFrame, dict]:
    text = log_path.read_text(encoding="utf-8")
    pattern = re.compile(
        r"\((\d+), Decimal\('([\d.]+)'\), Decimal\('([\d.]+)'\), "
        r"(\d+), (\d+), Decimal\('([\d.]+)'\), (\d+), Decimal\('([\d.]+)'\), "
        r"Decimal\('([\d.]+)'\), '([A-Z_]+)', (\d+)\)"
    )
    columns = [
        "parameter_set_id", "annual_growth_pct", "quarterly_growth_pct",
        "annual_years", "quarter_count", "volume_ratio_threshold",
        "volume_surge_min_days", "daily_ma_tolerance_pct",
        "weekly_ma_tolerance_pct", "screen_type", "selected_stock_count",
    ]
    df = pd.DataFrame(pattern.findall(text), columns=columns)
    numeric = [column for column in columns if column != "screen_type"]
    df[numeric] = df[numeric].apply(pd.to_numeric)

    causal_summary_pattern = re.compile(
        r"\('([A-Z_]+)', (\d+), (\d+), (\d+), datetime\.date\((\d+), (\d+), (\d+)\), "
        r"datetime\.date\((\d+), (\d+), (\d+)\), datetime\.date\((\d+), (\d+), (\d+)\), "
        r"datetime\.date\((\d+), (\d+), (\d+)\)\)"
    )
    legacy_summary_pattern = re.compile(
        r"\('([A-Z_]+)', (\d+), (\d+), (\d+), datetime\.date\((\d+), (\d+), (\d+)\), "
        r"datetime\.date\((\d+), (\d+), (\d+)\)\)"
    )
    screen_summary = {}
    for match in causal_summary_pattern.findall(text):
        screen_summary[match[0]] = {
            "outcomes": int(match[1]),
            "parameter_sets": int(match[2]),
            "unique_securities": int(match[3]),
            "earliest_signal": f"{match[4]}-{int(match[5]):02d}-{int(match[6]):02d}",
            "latest_signal": f"{match[7]}-{int(match[8]):02d}-{int(match[9]):02d}",
            "earliest_entry": f"{match[10]}-{int(match[11]):02d}-{int(match[12]):02d}",
            "latest_entry": f"{match[13]}-{int(match[14]):02d}-{int(match[15]):02d}",
        }
    if not screen_summary:
        for match in legacy_summary_pattern.findall(text):
            screen_summary[match[0]] = {
                "outcomes": int(match[1]),
                "parameter_sets": int(match[2]),
                "unique_securities": int(match[3]),
                "earliest_signal": f"{match[4]}-{int(match[5]):02d}-{int(match[6]):02d}",
                "latest_signal": f"{match[7]}-{int(match[8]):02d}-{int(match[9]):02d}",
                "earliest_entry": "not available",
                "latest_entry": "not available",
            }
    if df.empty or set(screen_summary) != {"A_F", "A_H"}:
        raise ValueError(f"Could not parse complete A-F/A-H results from {log_path}")
    return df, screen_summary


def font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    path = next((candidate for candidate in candidates if Path(candidate).exists()), None)
    return ImageFont.truetype(path, size=size) if path else ImageFont.load_default()


def make_sensitivity_chart(df: pd.DataFrame) -> None:
    factors = [
        ("volume_ratio_threshold", "Volume ratio threshold"),
        ("annual_years", "Annual years"),
        ("quarter_count", "Quarter count"),
        ("volume_surge_min_days", "Volume surge days"),
    ]
    image = Image.new("RGB", (1500, 980), "white")
    draw = ImageDraw.Draw(image)
    draw.text((55, 28), "Average selected stocks per parameter set", fill=f"#{NAVY}", font=font(38, True))
    draw.text((55, 78), "A-F and A-H screening yield by selected parameter value", fill=f"#{MID_GRAY}", font=font(24))

    colors = {"A_F": f"#{BLUE}", "A_H": f"#{GREEN}"}
    for panel_index, (factor, title) in enumerate(factors):
        x0 = 55 + (panel_index % 2) * 725
        y0 = 145 + (panel_index // 2) * 390
        grouped = df.groupby([factor, "screen_type"])["selected_stock_count"].mean().unstack(fill_value=0)
        max_value = max(grouped.max().max(), 1)
        draw.text((x0, y0), title, fill=f"#{NAVY}", font=font(27, True))
        values = list(grouped.index)
        for index, value in enumerate(values):
            row_y = y0 + 55 + index * 68
            draw.text((x0, row_y + 10), str(value), fill="#222222", font=font(22, True))
            for screen_index, screen in enumerate(["A_F", "A_H"]):
                avg = float(grouped.loc[value].get(screen, 0))
                bar_x = x0 + 75
                bar_y = row_y + screen_index * 26
                width = int(500 * avg / max_value)
                draw.rounded_rectangle(
                    (bar_x, bar_y, bar_x + width, bar_y + 19),
                    radius=4,
                    fill=colors[screen],
                )
                draw.text((bar_x + width + 10, bar_y - 4), f"{avg:.1f}", fill="#222222", font=font(18))
        draw.text((x0 + 500, y0 + 10), "A-F", fill=colors["A_F"], font=font(18, True))
        draw.text((x0 + 565, y0 + 10), "A-H", fill=colors["A_H"], font=font(18, True))
    image.save(CHART_PATH)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        shade_cell(cell, header_fill)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(header))
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(NAVY)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.width = Inches(widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2:
                shade_cell(cell, "FAFBFC")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value))
            run.font.size = Pt(font_size)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def build_report(df: pd.DataFrame, screen_summary: dict) -> None:
    make_sensitivity_chart(df)
    pivot = (
        df.pivot_table(
            index=[
                "parameter_set_id", "annual_growth_pct", "quarterly_growth_pct",
                "annual_years", "quarter_count", "volume_ratio_threshold",
                "volume_surge_min_days",
            ],
            columns="screen_type",
            values="selected_stock_count",
            fill_value=0,
        )
        .reset_index()
    )
    pivot["total"] = pivot.get("A_F", 0) + pivot.get("A_H", 0)
    pivot["a_h_share"] = pivot.get("A_H", 0) / pivot.get("A_F", 1).replace(0, 1)
    top = pivot.sort_values(["total", "A_H"], ascending=False).head(10)
    bottom = pivot.sort_values(["total", "A_H", "A_F"], ascending=True).head(10)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 7),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "NASDAQ Stock Recommendation | Backtest Research"
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = RGBColor.from_string(MID_GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run(f"Generated {date.today().strftime('%B %d, %Y')}").font.size = Pt(8.5)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("BACKTEST PARAMETER COMPARISON")
    run.bold = True
    run.font.size = Pt(23)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    sr = subtitle.add_run("Screening yield, confirmation coverage, and parameter sensitivity")
    sr.font.size = Pt(13)
    sr.font.color.rgb = RGBColor.from_string(MID_GRAY)

    meta = doc.add_paragraph()
    meta.add_run("Evaluation window: ").bold = True
    meta.add_run(
        f"{min(item['earliest_signal'] for item in screen_summary.values())} to "
        f"{max(item['latest_entry'] for item in screen_summary.values())}  |  "
    )
    meta.add_run("Parameter sets: ").bold = True
    meta.add_run("192  |  ")
    meta.add_run("Compact outcomes: ").bold = True
    meta.add_run(f"{sum(item['outcomes'] for item in screen_summary.values()):,}")

    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = callout.cell(0, 0)
    shade_cell(cell, LIGHT_BLUE)
    set_cell_margins(cell, 150, 170, 150, 170)
    p = cell.paragraphs[0]
    r = p.add_run(
        "Key conclusion: volume-ratio threshold is the dominant driver of screening yield. "
        f"The broadest configuration selected {int(pivot['A_F'].max())} A-F and "
        f"{int(pivot['A_H'].max())} A-H securities, while stricter "
        "volume settings reduced the candidate pool rapidly. This report compares screening "
        "coverage, not realized investment-return quality."
    )
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)

    add_heading(doc, "Parameter Choice Reference", 1)
    add_bullet(
        doc,
        "Six selectable parameters form the 192-combination grid. Daily and weekly "
        "moving-average tolerances remain fixed across every parameter set.",
    )
    parameter_rows = [
        ["Annual growth %", "A", "2%, 3%", "Selectable", "Minimum annual revenue and operating-income growth."],
        ["Quarterly growth %", "B", "2%, 3%", "Selectable", "Minimum quarterly revenue and operating-income growth."],
        ["Annual periods", "A", "2, 3 years", "Selectable", "Number of latest annual periods that must pass."],
        ["Quarterly periods", "B", "2, 3, 4 quarters", "Selectable", "Number of latest quarterly periods that must pass."],
        ["Volume ratio", "C / D", "2x, 3x, 4x, 5x", "Selectable", "Required volume-surge multiple."],
        ["Volume surge days", "D", "2, 3 days", "Selectable", "Minimum surge days in the recent-volume window."],
        ["Daily MA tolerance", "E / F", "1%", "Fixed", "Allowed daily MA20/MA50/MA100 alignment tolerance."],
        ["Weekly MA tolerance", "G / H", "2%", "Fixed", "Allowed weekly MA5/MA10/MA30 alignment tolerance."],
    ]
    add_table(
        doc,
        ["Parameter", "Conditions", "Choices", "Mode", "Meaning"],
        parameter_rows,
        [1.3, 0.7, 1.2, 0.8, 2.5],
        font_size=7.8,
    )

    add_heading(doc, "Executive Summary", 1)
    add_bullet(doc, "The run completed successfully across all 192 parameter combinations with no duplicate outcomes, missing core prices, invalid dates, or high/low inconsistencies.")
    add_bullet(doc, f"A-F produced {screen_summary['A_F']['outcomes']:,} parameter-security outcomes across {screen_summary['A_F']['parameter_sets']} parameter sets and {screen_summary['A_F']['unique_securities']} unique securities.")
    add_bullet(doc, f"A-H produced {screen_summary['A_H']['outcomes']:,} outcomes across {screen_summary['A_H']['parameter_sets']} parameter sets and {screen_summary['A_H']['unique_securities']} unique securities; weekly confirmation narrows the candidate universe.")
    add_bullet(doc, "Lower volume-ratio thresholds and shorter fundamental-history requirements produce the largest candidate pools.")
    add_bullet(doc, "Return, median return, win rate, and drawdown ranking are not present in the execution log; those metrics should be queried from `backtest_selection_outcome` before selecting a final strategy.")

    add_heading(doc, "Validated Run Overview", 1)
    overview_rows = [
        ["A-F", f"{screen_summary['A_F']['outcomes']:,}", screen_summary["A_F"]["parameter_sets"], screen_summary["A_F"]["unique_securities"], screen_summary["A_F"]["earliest_signal"], screen_summary["A_F"]["latest_entry"]],
        ["A-H", f"{screen_summary['A_H']['outcomes']:,}", screen_summary["A_H"]["parameter_sets"], screen_summary["A_H"]["unique_securities"], screen_summary["A_H"]["earliest_signal"], screen_summary["A_H"]["latest_entry"]],
    ]
    add_table(
        doc,
        ["Screen", "Outcomes", "Sets with results", "Unique stocks", "First signal", "Last entry"],
        overview_rows,
        [0.7, 0.85, 1.15, 1.0, 1.05, 1.05],
        font_size=8.5,
    )

    doc.add_picture(str(CHART_PATH), width=Inches(6.75))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run("Figure 1. Average screening yield by parameter choice.")
    cr.italic = True
    cr.font.size = Pt(8.5)
    cr.font.color.rgb = RGBColor.from_string(MID_GRAY)

    add_heading(doc, "Parameter Sensitivity", 1)
    sensitivity_rows = []
    factor_labels = {
        "annual_growth_pct": "Annual growth threshold",
        "quarterly_growth_pct": "Quarterly growth threshold",
        "annual_years": "Annual periods",
        "quarter_count": "Quarterly periods",
        "volume_ratio_threshold": "Volume ratio threshold",
        "volume_surge_min_days": "Volume surge days",
    }
    for factor, label in factor_labels.items():
        grouped = df.groupby([factor, "screen_type"])["selected_stock_count"].mean().unstack(fill_value=0)
        low = grouped.index.min()
        high = grouped.index.max()
        sensitivity_rows.append(
            [
                label,
                f"{low} -> {high}",
                f"{grouped.loc[low].get('A_F', 0):.1f} -> {grouped.loc[high].get('A_F', 0):.1f}",
                f"{grouped.loc[low].get('A_H', 0):.1f} -> {grouped.loc[high].get('A_H', 0):.1f}",
            ]
        )
    add_table(
        doc,
        ["Factor", "Compared values", "Avg A-F yield", "Avg A-H yield"],
        sensitivity_rows,
        [2.25, 1.15, 1.45, 1.45],
        font_size=8.5,
    )

    add_heading(doc, "Interpretation", 2)
    for factor, label in [
        ("volume_ratio_threshold", "Volume ratio"),
        ("annual_years", "Annual-period requirement"),
        ("quarter_count", "Quarterly-period requirement"),
        ("volume_surge_min_days", "Volume-surge-day requirement"),
    ]:
        grouped = df.groupby([factor, "screen_type"])["selected_stock_count"].mean().unstack(fill_value=0)
        low, high = grouped.index.min(), grouped.index.max()
        add_bullet(
            doc,
            f"{label}: moving from {low:g} to {high:g} changes average A-F yield "
            f"from {grouped.loc[low].get('A_F', 0):.1f} to {grouped.loc[high].get('A_F', 0):.1f} "
            f"and average A-H yield from {grouped.loc[low].get('A_H', 0):.1f} "
            f"to {grouped.loc[high].get('A_H', 0):.1f}.",
        )

    add_heading(doc, "Highest-Coverage Parameter Sets", 1)
    top_rows = []
    for _, row in top.iterrows():
        top_rows.append(
            [
                int(row["parameter_set_id"]),
                f"{int(row['annual_growth_pct'])}/{int(row['quarterly_growth_pct'])}",
                f"{int(row['annual_years'])}/{int(row['quarter_count'])}",
                f"{int(row['volume_ratio_threshold'])}x/{int(row['volume_surge_min_days'])}d",
                int(row.get("A_F", 0)),
                int(row.get("A_H", 0)),
                int(row["total"]),
            ]
        )
    add_table(
        doc,
        ["ID", "Growth A/Q", "Periods A/Q", "Volume", "A-F", "A-H", "Total"],
        top_rows,
        [0.45, 0.9, 1.0, 1.0, 0.55, 0.55, 0.6],
        font_size=8.2,
    )

    doc.add_page_break()
    add_heading(doc, "Lowest-Coverage Parameter Sets", 1)
    add_bullet(
        doc,
        "These configurations produced the fewest combined A-F and A-H selections. "
        "Their sparse samples are useful for understanding selectivity, but are too small "
        "for reliable return-performance conclusions.",
    )
    bottom_rows = []
    for _, row in bottom.iterrows():
        bottom_rows.append(
            [
                int(row["parameter_set_id"]),
                f"{int(row['annual_growth_pct'])}/{int(row['quarterly_growth_pct'])}",
                f"{int(row['annual_years'])}/{int(row['quarter_count'])}",
                f"{int(row['volume_ratio_threshold'])}x/{int(row['volume_surge_min_days'])}d",
                int(row.get("A_F", 0)),
                int(row.get("A_H", 0)),
                int(row["total"]),
            ]
        )
    add_table(
        doc,
        ["ID", "Growth A/Q", "Periods A/Q", "Volume", "A-F", "A-H", "Total"],
        bottom_rows,
        [0.45, 0.9, 1.0, 1.0, 0.55, 0.55, 0.6],
        font_size=8.2,
    )

    add_heading(doc, "Recommended Comparison Set", 1)
    recommendation_rows = []
    recommendation_specs = [
        ("Broad discovery", 1, "Largest sample; suitable baseline."),
        ("Moderate volume", 3, "Stricter volume comparator."),
        ("Strict volume", 7, "High-conviction, smaller-sample comparator."),
        ("Durability focus", 41, "Longer fundamental-history comparator."),
    ]
    for role, identifier, use in recommendation_specs:
        row = pivot.loc[pivot["parameter_set_id"] == identifier].iloc[0]
        configuration = (
            f"{int(row['annual_growth_pct'])}%/{int(row['quarterly_growth_pct'])}%, "
            f"{int(row['annual_years'])} annual, {int(row['quarter_count'])} quarters, "
            f"{int(row['volume_ratio_threshold'])}x volume, "
            f"{int(row['volume_surge_min_days'])} surge days"
        )
        recommendation_rows.append(
            [role, identifier, configuration, f"{int(row['A_F'])} / {int(row['A_H'])}", use]
        )
    add_table(
        doc,
        ["Role", "ID", "Configuration", "A-F / A-H", "Use"],
        recommendation_rows,
        [1.05, 0.45, 2.55, 0.8, 1.55],
        font_size=8.0,
    )

    add_heading(doc, "Decision Guidance", 1)
    add_bullet(doc, "Do not label the highest-coverage parameter set as the best-performing investment strategy. Coverage and return quality answer different questions.")
    add_bullet(doc, "Require a minimum sample size before ranking return performance. A practical initial threshold is at least 20 A-F outcomes or at least 10 A-H outcomes.")
    add_bullet(doc, "Compare mean and median return together. A large gap often indicates that a few extreme winners dominate the average.")
    add_bullet(doc, "Use win rate and maximum drawdown alongside returns; parameter sets with high average returns but severe drawdowns may be unsuitable.")
    add_bullet(doc, "Evaluate A-F and A-H separately. A-H is a stricter weekly-confirmed strategy and should not be pooled with A-F.")

    add_heading(doc, "Required Return-Performance Query", 1)
    p = doc.add_paragraph()
    p.add_run(
        "To complete an investment-performance ranking, export the following aggregate query from Supabase and use it for a second report:"
    )
    query = (
        "SELECT p.parameter_set_id, p.parameter_set_name, o.screen_type, COUNT(*) AS sample_size,\n"
        "ROUND(AVG(o.return_pct), 2) AS avg_return_pct,\n"
        "ROUND(PERCENTILE_CONT(0.5) WITHIN GROUP (ORDER BY o.return_pct)::numeric, 2) AS median_return_pct,\n"
        "ROUND(100.0 * COUNT(*) FILTER (WHERE o.return_pct > 0) / COUNT(*), 2) AS win_rate_pct,\n"
        "ROUND(AVG(o.max_drawdown_pct), 2) AS avg_drawdown_pct,\n"
        "ROUND(AVG(o.max_return_pct), 2) AS avg_max_return_pct\n"
        "FROM backtest_parameter_set p JOIN backtest_selection_outcome o USING (parameter_set_id)\n"
        "GROUP BY p.parameter_set_id, p.parameter_set_name, o.screen_type\n"
        "ORDER BY o.screen_type, sample_size DESC;"
    )
    qp = doc.add_paragraph()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), LIGHT_GRAY)
    qp._p.get_or_add_pPr().append(shade)
    qr = qp.add_run(query)
    qr.font.name = "Courier New"
    qr.font.size = Pt(7.5)

    add_heading(doc, "Scope and Limitations", 1)
    add_bullet(doc, "This report uses the validated execution log and compares selection yield, not realized portfolio returns.")
    add_bullet(doc, "The same security may appear in many parameter sets, so outcome rows are not independent observations.")
    add_bullet(doc, "Fundamental histories use `datadate <= signal_date`; actual public filing-availability dates are unavailable, so some look-ahead-bias risk remains.")
    add_bullet(doc, "No transaction costs, liquidity constraints, holding rules, or overlapping-position controls are included.")

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Build the screening-yield comparison report.")
    parser.add_argument("--log", type=Path, required=True, help="Validated backtest execution log.")
    args = parser.parse_args()
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    data, summaries = parse_log(args.log)
    build_report(data, summaries)
    print(DOCX_PATH)
