from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPORT_DIR = Path(__file__).resolve().parent
OUTPUT = REPORT_DIR / "backtest_plan_and_flow_report.docx"

INK = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(92, 107, 128)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "E8F3EC"
PALE_GOLD = "FFF4CE"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths, first_col_left=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Inches(widths[index])
        shade_cell(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if first_col_left and index == 0 else WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = INK
    for row_index, values in enumerate(rows):
        table_row = table.add_row()
        table_row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for index, value in enumerate(values):
            cell = table_row.cells[index]
            cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index % 2:
                shade_cell(cell, LIGHT_GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if first_col_left and index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(str(value)).font.size = Pt(8.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    return paragraph


def add_numbered(doc, title, detail):
    paragraph = doc.add_paragraph(style="List Number")
    title_run = paragraph.add_run(f"{title}: ")
    title_run.bold = True
    paragraph.add_run(detail)
    return paragraph


def add_callout(doc, title, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, 150, 180, 150, 180)
    paragraph = cell.paragraphs[0]
    title_run = paragraph.add_run(f"{title}: ")
    title_run.bold = True
    title_run.font.color.rgb = INK
    paragraph.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for list_name in ["List Bullet", "List Number"]:
        style = doc.styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.1
    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11, INK),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
    header = section.header.paragraphs[0]
    header.add_run("NASDAQ Stock Recommendation | Backtest Design").font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run(f"Verified {date.today().strftime('%B %d, %Y')}").font.color.rgb = MUTED


def build_report() -> Path:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    run = title.add_run("BACKTEST PLAN AND VERIFIED EXECUTION FLOW")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = INK
    subtitle = doc.add_paragraph()
    subtitle.add_run(
        "How 192 screening parameter sets are evaluated from January 2022 through fixed return horizons"
    ).font.color.rgb = MUTED

    add_callout(
        doc,
        "Verdict",
        "The proposed flow is correct, with one important detail: for each parameter set, "
        "screen type, and security, the implementation retains only the earliest qualifying "
        "selection that completes all required confirmations.",
        PALE_GREEN,
    )

    add_heading(doc, "Verified Flow")
    add_numbered(
        doc,
        "Create 192 parameter sets",
        "Build the Cartesian product of six selectable screening parameters. Daily and weekly moving-average tolerances remain fixed.",
    )
    add_numbered(
        doc,
        "Build historical features",
        "Read production S3 daily price data and fundamental growth histories, then build daily and completed-week feature Parquet files with DuckDB.",
    )
    add_numbered(
        doc,
        "Find A-F and A-H selections",
        "Evaluate every parameter set from January 1, 2022. A-F becomes actionable after the next-day F confirmation; A-H becomes actionable after G and the following completed-week H confirmation.",
    )
    add_numbered(
        doc,
        "Keep the earliest selection per security",
        "Within each parameter set and screen type, later qualifying signals for the same gvkey/iid security are discarded.",
    )
    add_numbered(
        doc,
        "Calculate fixed-horizon returns",
        "Measure each retained selection at 6 months, 1 year, and 2 years using the first available trading price on or after the calendar horizon.",
    )
    add_numbered(
        doc,
        "Compare parameter-set performance",
        "For each screen and horizon, aggregate completed sample count, average return, median return, and win rate; then combine the three average-return ranks.",
    )

    add_heading(doc, "Parameter Grid")
    add_table(
        doc,
        ["Parameter", "Choices", "Condition"],
        [
            ["Annual growth threshold", "2%, 3%", "A"],
            ["Quarterly growth threshold", "2%, 3%", "B"],
            ["Annual periods", "2, 3 years", "A"],
            ["Quarterly periods", "2, 3, 4 quarters", "B"],
            ["Volume-ratio threshold", "2x, 3x, 4x, 5x", "C / D"],
            ["Volume-surge minimum days", "2, 3 days", "D"],
            ["Daily MA tolerance", "1% fixed", "E / F"],
            ["Weekly MA tolerance", "2% fixed", "G / H"],
        ],
        [2.25, 2.4, 1.2],
    )
    doc.add_paragraph(
        "Combination count: 2 x 2 x 2 x 3 x 4 x 2 = 192 parameter sets."
    )

    add_heading(doc, "Causal Selection Timing")
    add_callout(
        doc,
        "No pre-confirmation return",
        "Entry price and return measurement begin only when the required confirmation is observable. "
        "Price movement before the actionable selected date is excluded.",
    )
    add_table(
        doc,
        ["Stage", "A-F", "A-H"],
        [
            ["A-E signal", "Daily signal date", "Daily signal date"],
            ["F confirmation", "Next trading row", "Next trading row"],
            ["G confirmation", "Not required", "First completed official week on/after F"],
            ["H confirmation", "Not required", "Following completed official week"],
            ["Actionable selected date", "F confirmation date", "H confirmation date"],
            ["Entry price", "F-confirmation adjusted close fallback", "H-confirmation weekly close"],
        ],
        [1.8, 2.15, 2.35],
    )

    add_heading(doc, "Selection Unit")
    doc.add_paragraph(
        "The stored observation is not every signal event. It is the earliest qualifying security selection for:"
    )
    add_bullet(doc, "one parameter set")
    add_bullet(doc, "one screen type: A-F or A-H")
    add_bullet(doc, "one security identified by gvkey and iid")
    add_callout(
        doc,
        "Consequence",
        "A stock that qualifies again later under the same parameter set and screen does not create another outcome. "
        "This avoids repeated-event concentration, but it means the study evaluates first-selection performance rather than every trading opportunity.",
        PALE_GOLD,
    )

    add_heading(doc, "Fixed-Horizon Return Calculation")
    add_table(
        doc,
        ["Horizon", "Price date used", "Return formula"],
        [
            ["6 months", "First trading date on/after selected date + 6 calendar months", "(price_6m / entry price - 1) x 100"],
            ["1 year", "First trading date on/after selected date + 1 calendar year", "(price_1y / entry price - 1) x 100"],
            ["2 years", "First trading date on/after selected date + 2 calendar years", "(price_2y / entry price - 1) x 100"],
        ],
        [1.0, 3.5, 2.0],
    )
    doc.add_paragraph(
        "If a selection has not yet reached a horizon, its horizon return remains null and is excluded from that horizon's comparison."
    )

    add_heading(doc, "Performance Comparison")
    add_bullet(doc, "A-F and A-H are compared separately because they have different confirmation timing and entry dates.")
    add_bullet(doc, "Each horizon is ranked independently by average return after applying a minimum completed-sample threshold.")
    add_bullet(doc, "The final combined score is the mean of the 6-month, 1-year, and 2-year average-return ranks; lower is better.")
    add_bullet(doc, "Combined eligibility requires sufficient completed observations at all three horizons.")

    add_heading(doc, "Latest Verified Run")
    add_table(
        doc,
        ["Screen", "Selections", "Unique securities", "6m complete", "1y complete", "2y complete"],
        [
            ["A-F", "3,698", "92", "3,364", "3,050", "2,436"],
            ["A-H", "1,928", "47", "1,748", "1,576", "1,292"],
        ],
        [0.75, 1.0, 1.2, 1.0, 1.0, 1.0],
    )
    doc.add_paragraph(
        "All 192 parameter sets produced stored outcomes for both screens. Validation reported zero missing core price outcomes, "
        "zero confirmation-timing errors, zero inconsistent fixed-horizon rows, and zero invalid horizon dates."
    )

    add_heading(doc, "What This Backtest Answers")
    add_bullet(doc, "Which parameter sets historically selected securities with stronger first-selection returns at each fixed horizon?")
    add_bullet(doc, "Which parameter sets performed consistently across 6-month, 1-year, and 2-year comparisons?")
    add_bullet(doc, "How much completed sample coverage supports each ranking?")

    add_heading(doc, "What It Does Not Yet Prove")
    add_bullet(doc, "A higher rank does not prove statistically significant superiority because parameter sets often reuse the same securities and dates.")
    add_bullet(doc, "The current design does not simulate repeated entries, portfolio sizing, transaction costs, or overlapping-position capital constraints.")
    add_bullet(doc, "Fundamentals use datadate <= signal date because filing/publication availability dates are not present; this reduces but does not fully remove look-ahead risk.")
    add_bullet(doc, "Fixed-horizon outcomes describe hypothetical holding periods, not an implemented exit strategy.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_report())
