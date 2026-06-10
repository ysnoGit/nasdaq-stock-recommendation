from __future__ import annotations

import argparse
import json
import statistics
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPORT_DIR = Path(__file__).resolve().parent
DEFAULT_INPUT = REPORT_DIR / "performance_comparison.json"
DEFAULT_OUTPUT = REPORT_DIR / "backtest_performance_comparison_report.docx"
EXPECTED_SETS_PER_SCREEN = 192
HORIZONS = [
    ("6m", "6 months"),
    ("1y", "1 year"),
    ("2y", "2 years"),
]
MINIMUMS = {"A_F": 30, "A_H": 15}

INK = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(92, 107, 128)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "E8F3EC"


def load_rows(path: Path) -> list[dict]:
    rows = json.loads(path.read_text(encoding="utf-8"))
    integer_fields = [
        "parameter_set_id",
        "sample_size",
        "sample_size_6m",
        "sample_size_1y",
        "sample_size_2y",
    ]
    numeric_fields = [
        "annual_growth_pct",
        "quarterly_growth_pct",
        "annual_years",
        "quarter_count",
        "volume_ratio_threshold",
        "volume_surge_min_days",
    ]
    for horizon, _ in HORIZONS:
        numeric_fields.extend(
            [
                f"avg_return_{horizon}_pct",
                f"median_return_{horizon}_pct",
                f"win_rate_{horizon}_pct",
            ]
        )
    for row in rows:
        for key in integer_fields:
            row[key] = int(row[key])
        for key in numeric_fields:
            row[key] = None if row[key] is None else float(row[key])

    counts = {
        screen: sum(row["screen_type"] == screen for row in rows)
        for screen in {"A_F", "A_H"}
    }
    expected = {"A_F": EXPECTED_SETS_PER_SCREEN, "A_H": EXPECTED_SETS_PER_SCREEN}
    if counts != expected:
        raise ValueError(f"Expected 192 complete parameter sets for each screen; found {counts}")
    return rows


def rank_horizon(rows: list[dict], horizon: str, minimum: int) -> list[dict]:
    sample_key = f"sample_size_{horizon}"
    return_key = f"avg_return_{horizon}_pct"
    eligible = [
        dict(row)
        for row in rows
        if row[sample_key] >= minimum and row[return_key] is not None
    ]
    return sorted(eligible, key=lambda row: (-row[return_key], -row[sample_key], row["parameter_set_id"]))


def combined_ranking(rows: list[dict], minimum: int) -> list[dict]:
    eligible = [
        dict(row)
        for row in rows
        if all(
            row[f"sample_size_{horizon}"] >= minimum
            and row[f"avg_return_{horizon}_pct"] is not None
            for horizon, _ in HORIZONS
        )
    ]
    by_id = {row["parameter_set_id"]: row for row in eligible}
    for horizon, _ in HORIZONS:
        ordered = sorted(
            eligible,
            key=lambda row: (
                -row[f"avg_return_{horizon}_pct"],
                -row[f"sample_size_{horizon}"],
                row["parameter_set_id"],
            ),
        )
        for rank, row in enumerate(ordered, 1):
            by_id[row["parameter_set_id"]][f"rank_{horizon}"] = rank
    for row in eligible:
        row["combined_rank_score"] = statistics.mean(
            row[f"rank_{horizon}"] for horizon, _ in HORIZONS
        )
    return sorted(
        eligible,
        key=lambda row: (row["combined_rank_score"], row["parameter_set_id"]),
    )


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths, highlight_first=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Inches(widths[index])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(header))
        run.bold = True
        run.font.size = Pt(8.2)
        run.font.color.rgb = INK
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if highlight_first and row_index == 0:
                shade_cell(cell, PALE_GREEN)
            elif row_index % 2:
                shade_cell(cell, LIGHT_GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 1 else WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(str(value)).font.size = Pt(8.0)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    return paragraph


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11, INK),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
    header = section.header.paragraphs[0]
    header.add_run("NASDAQ Stock Recommendation | Fixed-Horizon Backtest").font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run(f"Generated {date.today().strftime('%B %d, %Y')}").font.color.rgb = MUTED


def pct(value) -> str:
    return "n/a" if value is None else f"{value:.2f}%"


def combined_row(rank: int, row: dict) -> list:
    return [
        rank,
        row["parameter_set_name"],
        f"{row['sample_size_6m']}/{row['sample_size_1y']}/{row['sample_size_2y']}",
        pct(row["avg_return_6m_pct"]),
        pct(row["avg_return_1y_pct"]),
        pct(row["avg_return_2y_pct"]),
        f"{row['rank_6m']}/{row['rank_1y']}/{row['rank_2y']}",
        f"{row['combined_rank_score']:.2f}",
    ]


def horizon_row(rank: int, row: dict, horizon: str) -> list:
    return [
        rank,
        row["parameter_set_name"],
        row[f"sample_size_{horizon}"],
        pct(row[f"avg_return_{horizon}_pct"]),
        pct(row[f"median_return_{horizon}_pct"]),
        pct(row[f"win_rate_{horizon}_pct"]),
    ]


def build_report(rows: list[dict], output: Path) -> Path:
    by_screen = {
        screen: [row for row in rows if row["screen_type"] == screen]
        for screen in ["A_F", "A_H"]
    }
    horizon_rankings = {
        screen: {
            horizon: rank_horizon(screen_rows, horizon, MINIMUMS[screen])
            for horizon, _ in HORIZONS
        }
        for screen, screen_rows in by_screen.items()
    }
    combined = {
        screen: combined_ranking(screen_rows, MINIMUMS[screen])
        for screen, screen_rows in by_screen.items()
    }

    doc = Document()
    configure_document(doc)
    title = doc.add_paragraph()
    run = title.add_run("FIXED-HORIZON PERFORMANCE COMPARISON")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = INK
    subtitle = doc.add_paragraph()
    subtitle.add_run(
        "Parameter-set results at 6 months, 1 year, and 2 years after actionable entry"
    ).font.color.rgb = MUTED

    lead = doc.add_table(rows=1, cols=1)
    shade_cell(lead.cell(0, 0), LIGHT_BLUE)
    set_cell_margins(lead.cell(0, 0), 150, 170, 150, 170)
    lead.cell(0, 0).paragraphs[0].add_run(
        "Each return uses the first available trading price on or after the calendar horizon. "
        "Unmatured selections are excluded. Final combined rank is the mean of the three "
        "average-return ranks and requires the minimum completed sample at every horizon."
    ).bold = True

    add_heading(doc, "Method")
    add_bullet(doc, "6-month return: price on the first trading date on or after entry date plus 6 calendar months.")
    add_bullet(doc, "1-year and 2-year returns follow the same rule using calendar-year anniversaries.")
    add_bullet(doc, "Average, median, and win rate use only selections that have completed that horizon.")
    add_bullet(doc, "A-F and A-H are ranked separately because their confirmation timing and entry dates differ.")

    add_heading(doc, "Screen Overview")
    overview = []
    for screen, screen_rows in by_screen.items():
        overview.append(
            [
                screen.replace("_", "-"),
                len(screen_rows),
                f"{sum(row['sample_size'] for row in screen_rows):,}",
                f"{sum(row['sample_size_6m'] for row in screen_rows):,}",
                f"{sum(row['sample_size_1y'] for row in screen_rows):,}",
                f"{sum(row['sample_size_2y'] for row in screen_rows):,}",
                MINIMUMS[screen],
                len(combined[screen]),
            ]
        )
    add_table(
        doc,
        ["Screen", "Sets", "Selections", "6m complete", "1y complete", "2y complete", "Min N", "Combined eligible"],
        overview,
        [0.55, 0.45, 0.75, 0.85, 0.85, 0.85, 0.55, 1.0],
    )

    add_heading(doc, "Combined Leaders")
    for screen in ["A_F", "A_H"]:
        if combined[screen]:
            leader = combined[screen][0]
            add_bullet(
                doc,
                f"{screen.replace('_', '-')}: {leader['parameter_set_name']} leads with "
                f"{pct(leader['avg_return_6m_pct'])}, {pct(leader['avg_return_1y_pct'])}, "
                f"and {pct(leader['avg_return_2y_pct'])} average returns at 6m/1y/2y; "
                f"combined rank score {leader['combined_rank_score']:.2f}.",
            )
        else:
            add_bullet(doc, f"{screen.replace('_', '-')}: no set meets the completed-sample threshold at all horizons.")

    for screen in ["A_F", "A_H"]:
        doc.add_page_break()
        label = screen.replace("_", "-")
        add_heading(doc, f"{label} Combined Ranking")
        doc.add_paragraph(
            f"Minimum completed observations per horizon: {MINIMUMS[screen]}. "
            "Lower combined rank score is better. N is shown as 6m/1y/2y."
        )
        add_table(
            doc,
            ["Rank", "Parameter set", "N 6m/1y/2y", "Avg 6m", "Avg 1y", "Avg 2y", "Ranks 6m/1y/2y", "Combined"],
            [combined_row(index, row) for index, row in enumerate(combined[screen][:15], 1)],
            [0.4, 1.65, 0.85, 0.7, 0.7, 0.7, 1.0, 0.65],
            highlight_first=True,
        )

        for horizon, horizon_label in HORIZONS:
            add_heading(doc, f"{label} {horizon_label} Ranking", 2)
            doc.add_paragraph(
                f"Ranked by average {horizon_label} return among sets with at least "
                f"{MINIMUMS[screen]} completed observations."
            )
            add_table(
                doc,
                ["Rank", "Parameter set", "N", "Average", "Median", "Win rate"],
                [
                    horizon_row(index, row, horizon)
                    for index, row in enumerate(horizon_rankings[screen][horizon][:12], 1)
                ],
                [0.45, 2.1, 0.55, 0.85, 0.85, 0.85],
                highlight_first=True,
            )

    add_heading(doc, "Interpretation and Limitations")
    add_bullet(doc, "The combined rank rewards consistency across horizons; it does not weight horizons by economic importance.")
    add_bullet(doc, "Completed-horizon filtering means newer selections contribute to 6-month results before they can contribute to 1-year or 2-year results.")
    add_bullet(doc, "Returns do not yet model transaction costs, dividends not captured by adjusted prices, portfolio overlap, position sizing, or delisting bias.")
    add_bullet(doc, "Parameter sets reuse many securities and dates, so rank differences are descriptive rather than proof of statistical significance.")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return output


def main() -> None:
    parser = argparse.ArgumentParser(description="Build the fixed-horizon performance comparison report.")
    parser.add_argument("--input", type=Path, default=DEFAULT_INPUT)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    print(build_report(load_rows(args.input), args.output))


if __name__ == "__main__":
    main()
